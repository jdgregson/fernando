"""Document generation helpers for PDF and DOCX from markdown-like content."""

import re
import os


def _parse_lines(content):
    """Parse markdown-like content into structured blocks."""
    blocks = []
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Code block
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            blocks.append(("code", "\n".join(code_lines)))
            i += 1
            continue

        # Image
        img_m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line.strip())
        if img_m:
            blocks.append(("image", img_m.group(2), img_m.group(1)))
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,3})\s+(.*)", line)
        if m:
            blocks.append(("heading", len(m.group(1)), m.group(2).strip()))
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^-{3,}\s*$", line.strip()):
            blocks.append(("hr",))
            i += 1
            continue

        # Table
        if "|" in line and i + 1 < len(lines) and re.match(r"^\|?[\s\-:|]+\|", lines[i + 1]):
            rows = []
            while i < len(lines) and "|" in lines[i]:
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not re.match(r"^[\s\-:|]+$", lines[i].strip().strip("|")):
                    rows.append(cells)
                i += 1
            blocks.append(("table", rows))
            continue

        # Bullet list
        if re.match(r"^\s*[-*]\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                items.append(re.sub(r"^\s*[-*]\s+", "", lines[i]))
                i += 1
            blocks.append(("bullet", items))
            continue

        # Numbered list
        if re.match(r"^\s*\d+\.\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                items.append(re.sub(r"^\s*\d+\.\s+", "", lines[i]))
                i += 1
            blocks.append(("number", items))
            continue

        # Paragraph (skip blank lines)
        if line.strip():
            blocks.append(("para", line))
        i += 1

    return blocks


def _inline_parts(text):
    """Split text into (style, text) parts for bold/italic/code."""
    parts = []
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            parts.append(("normal", text[last:m.start()]))
        if m.group(2):
            parts.append(("bold", m.group(2)))
        elif m.group(3):
            parts.append(("italic", m.group(3)))
        elif m.group(4):
            parts.append(("code", m.group(4)))
        last = m.end()
    if last < len(text):
        parts.append(("normal", text[last:]))
    return parts


def create_pdf(path, content, title=None):
    """Create a PDF from markdown-like content using fpdf2."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Unicode fonts
    FONTS_DIR = "/usr/share/fonts/truetype/dejavu"
    pdf.add_font("DejaVu", "", os.path.join(FONTS_DIR, "DejaVuSans.ttf"))
    pdf.add_font("DejaVu", "B", os.path.join(FONTS_DIR, "DejaVuSans-Bold.ttf"))
    pdf.add_font("DejaVu", "I", os.path.join(FONTS_DIR, "DejaVuSerif.ttf"))  # Serif as italic stand-in
    pdf.add_font("DejaVuMono", "", os.path.join(FONTS_DIR, "DejaVuSansMono.ttf"))
    pdf.add_font("DejaVuMono", "B", os.path.join(FONTS_DIR, "DejaVuSansMono-Bold.ttf"))
    FONT = "DejaVu"
    MONO = "DejaVuMono"

    if title:
        pdf.set_font(FONT, "B", 18)
        pdf.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

    def write_inline(text, size=11):
        for style, t in _inline_parts(text):
            if style == "bold":
                pdf.set_font(FONT, "B", size)
            elif style == "italic":
                pdf.set_font(FONT, "I", size)
            elif style == "code":
                pdf.set_font(MONO, "", size - 1)
            else:
                pdf.set_font(FONT, "", size)
            pdf.write(6, t)

    for block in _parse_lines(content):
        kind = block[0]

        if kind == "heading":
            level, text = block[1], block[2]
            sizes = {1: 16, 2: 14, 3: 12}
            pdf.ln(4)
            pdf.set_font(FONT, "B", sizes.get(level, 12))
            pdf.cell(0, 10, text, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)

        elif kind == "para":
            pdf.set_font(FONT, "", 11)
            write_inline(block[1])
            pdf.ln(8)

        elif kind == "bullet":
            for item in block[1]:
                pdf.set_font(FONT, "", 11)
                pdf.cell(8)
                pdf.write(6, "\u2022  ")
                write_inline(item)
                pdf.ln(7)
            pdf.ln(2)

        elif kind == "number":
            for idx, item in enumerate(block[1], 1):
                pdf.set_font(FONT, "", 11)
                pdf.cell(8)
                pdf.write(6, f"{idx}.  ")
                write_inline(item)
                pdf.ln(7)
            pdf.ln(2)

        elif kind == "code":
            pdf.set_font(MONO, "", 9)
            pdf.set_fill_color(240, 240, 240)
            for line in block[1].split("\n"):
                pdf.cell(0, 5, "  " + line, new_x="LMARGIN", new_y="NEXT", fill=True)
            pdf.ln(4)

        elif kind == "hr":
            y = pdf.get_y()
            pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
            pdf.ln(6)

        elif kind == "image":
            img_path = block[1]
            if os.path.isfile(img_path):
                max_w = pdf.w - pdf.l_margin - pdf.r_margin
                pdf.image(img_path, w=max_w)
                pdf.ln(4)

        elif kind == "table":
            rows = block[1]
            if not rows:
                continue
            n_cols = len(rows[0])
            col_w = (pdf.w - pdf.l_margin - pdf.r_margin) / n_cols
            # Header
            pdf.set_font(FONT, "B", 10)
            pdf.set_fill_color(230, 230, 230)
            for cell in rows[0]:
                pdf.cell(col_w, 7, cell, border=1, fill=True)
            pdf.ln()
            # Data rows
            pdf.set_font(FONT, "", 10)
            for row in rows[1:]:
                for j, cell in enumerate(row):
                    pdf.cell(col_w, 7, cell if j < len(row) else "", border=1)
                pdf.ln()
            pdf.ln(4)

    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    pdf.output(path)
    return path


def create_docx(path, content, title=None):
    """Create a DOCX from markdown-like content using python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Style defaults
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if title:
        p = doc.add_heading(title, level=0)

    def add_inline(paragraph, text):
        for style, t in _inline_parts(text):
            run = paragraph.add_run(t)
            if style == "bold":
                run.bold = True
            elif style == "italic":
                run.italic = True
            elif style == "code":
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for block in _parse_lines(content):
        kind = block[0]

        if kind == "heading":
            doc.add_heading(block[2], level=block[1])

        elif kind == "para":
            p = doc.add_paragraph()
            add_inline(p, block[1])

        elif kind == "bullet":
            for item in block[1]:
                p = doc.add_paragraph(style="List Bullet")
                add_inline(p, item)

        elif kind == "number":
            for item in block[1]:
                p = doc.add_paragraph(style="List Number")
                add_inline(p, item)

        elif kind == "code":
            p = doc.add_paragraph()
            run = p.add_run(block[1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Pt(18)

        elif kind == "hr":
            p = doc.add_paragraph()
            p.add_run("_" * 60).font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

        elif kind == "image":
            img_path = block[1]
            if os.path.isfile(img_path):
                from docx.shared import Inches
                doc.add_picture(img_path, width=Inches(6))

        elif kind == "table":
            rows = block[1]
            if not rows:
                continue
            table = doc.add_table(rows=len(rows), cols=len(rows[0]), style="Table Grid")
            for i, row in enumerate(rows):
                for j, cell in enumerate(row):
                    table.rows[i].cells[j].text = cell
                    if i == 0:
                        for run in table.rows[i].cells[j].paragraphs[0].runs:
                            run.bold = True

    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    doc.save(path)
    return path
